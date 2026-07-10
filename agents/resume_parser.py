# agents/resume_parser.py
# The Resume Parser Agent for the Korgut Commons.
# Reads a student resume (PDF or DOCX) and extracts a structured profile for Aria.

from __future__ import annotations

import base64
import json
import os
from pathlib import Path
from typing import Any, Dict, List

import anthropic
from rich.console import Console

console = Console()

MODEL = "claude-haiku-4-5-20251001"

EXTRACTION_PROMPT = """
You are a resume parser for Korgut, a graduate admissions platform.
Extract structured information from this student resume.

Return ONLY a valid JSON object with these exact fields.
No markdown, no explanation, just the JSON.

Required fields (use null if not found, never invent data):
{
  "name": string or null,
  "email": string or null,
  "undergraduate_institution": string or null,
  "undergraduate_major": string or null,
  "graduation_year": integer or null,
  "gpa": float or null,
  "gpa_scale": "4.0" or "10.0" or "percentage" or null,
  "gre_quant": integer or null,
  "gre_verbal": integer or null,
  "toefl": integer or null,
  "ielts": float or null,
  "work_experience_months": integer,
  "work_experience_summary": string or null,
  "research_experience": string or null,
  "publications_count": integer,
  "technical_skills": [list of strings],
  "projects": [
    {"title": string, "description": string, "technologies": [strings]}
  ],
  "inferred_disciplines": [list of strings],
  "gaps": [list of fields that are missing but important],
  "confidence_notes": string
}

Rules:
- Never invent data. If a field is not clearly present, use null.
- For work_experience_months: calculate from dates. Use 0 if no experience.
- For inferred_disciplines: suggest 2-3 graduate disciplines based on
  the student's major, skills, and projects.
- For gaps: always include 'budget' and 'target_disciplines'.
- For confidence_notes: one sentence about anything ambiguous or notable.
"""


def _get_anthropic_client() -> anthropic.Anthropic:
    """Create Anthropic client only when parsing is actually requested."""
    if not os.getenv("ANTHROPIC_API_KEY"):
        raise RuntimeError(
            "ANTHROPIC_API_KEY not found. Add it to your .env file before parsing resumes."
        )

    return anthropic.Anthropic()


def read_pdf(file_path: str) -> Dict[str, Any]:
    """
    Read PDF as a base64 document for Claude.

    Raises clear errors for missing/unreadable files instead of failing silently.
    """
    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"Resume PDF not found: {path}")

    if not path.is_file():
        raise ValueError(f"Resume path is not a file: {path}")

    try:
        with open(path, "rb") as file:
            data = base64.standard_b64encode(file.read()).decode("utf-8")
    except Exception as exc:
        raise RuntimeError(f"Failed to read PDF resume: {path}. Details: {exc}") from exc

    return {
        "type": "document",
        "source": {
            "type": "base64",
            "media_type": "application/pdf",
            "data": data,
        },
    }


def read_docx(file_path: str) -> str:
    """
    Read DOCX resume text safely.

    Note: legacy .doc files are not supported by python-docx.
    Convert .doc to .docx first if needed.
    """
    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"Resume DOCX not found: {path}")

    if not path.is_file():
        raise ValueError(f"Resume path is not a file: {path}")

    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError(
            "python-docx is not installed. Install it using: pip install python-docx"
        ) from exc

    try:
        document = Document(path)
        text_parts: List[str] = []

        for paragraph in document.paragraphs:
            if paragraph.text and paragraph.text.strip():
                text_parts.append(paragraph.text.strip())

        # Also read tables because resumes often keep education/skills in tables.
        for table in document.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = " ".join(
                        paragraph.text.strip()
                        for paragraph in cell.paragraphs
                        if paragraph.text and paragraph.text.strip()
                    )
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    text_parts.append(" | ".join(row_text))

        return "\n".join(text_parts)

    except Exception as exc:
        raise RuntimeError(f"Failed to read DOCX resume: {path}. Details: {exc}") from exc


def read_text_file(file_path: str) -> str:
    """Optional helper for plain text resumes or copied resume dumps."""
    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"Resume text file not found: {path}")

    if not path.is_file():
        raise ValueError(f"Resume path is not a file: {path}")

    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        return path.read_text(encoding="latin-1")


class ResumeParserAgent:
    """Parse PDF/DOCX/TXT resumes into the internal Korgut student profile schema."""

    def parse(self, file_path: str) -> Dict[str, Any]:
        path = Path(file_path)
        suffix = path.suffix.lower()

        console.print(f"[dim]Parsing: {path.name}[/dim]")

        if not path.exists():
            raise FileNotFoundError(f"Resume not found: {path}")

        if suffix == ".pdf":
            return self._parse_pdf(str(path))

        if suffix == ".docx":
            return self._parse_docx(str(path))

        if suffix == ".txt":
            return self._parse_text(str(path))

        if suffix == ".doc":
            raise ValueError(
                "Legacy .doc resumes are not supported directly. Convert the file to .docx or PDF first."
            )

        raise ValueError(
            f"Unsupported resume format: {suffix}. Supported formats: .pdf, .docx, .txt"
        )

    def _parse_pdf(self, file_path: str) -> Dict[str, Any]:
        """Extract structured profile from PDF using Claude document input."""
        try:
            document_content = read_pdf(file_path)
            client = _get_anthropic_client()

            response = client.messages.create(
                model=MODEL,
                max_tokens=1500,
                messages=[
                    {
                        "role": "user",
                        "content": [
                            document_content,
                            {"type": "text", "text": EXTRACTION_PROMPT},
                        ],
                    }
                ],
            )

            if not response.content:
                raise RuntimeError("Claude returned an empty response.")

            return self._process(response.content[0].text)

        except Exception as exc:
            raise RuntimeError(
                "Claude resume extraction failed for PDF. "
                "Check ANTHROPIC_API_KEY, model name, credits, network, and resume file. "
                f"Details: {exc}"
            ) from exc

    def _parse_docx(self, file_path: str) -> Dict[str, Any]:
        """Extract structured profile from DOCX using Claude."""
        text = read_docx(file_path)
        return self._parse_resume_text(text, source_type="DOCX")

    def _parse_text(self, file_path: str) -> Dict[str, Any]:
        """Extract structured profile from a plain text resume."""
        text = read_text_file(file_path)
        return self._parse_resume_text(text, source_type="TXT")

    def _parse_resume_text(self, text: str, source_type: str = "text") -> Dict[str, Any]:
        if not text or not text.strip():
            raise RuntimeError(f"{source_type} resume text is empty.")

        try:
            client = _get_anthropic_client()

            response = client.messages.create(
                model=MODEL,
                max_tokens=1500,
                messages=[
                    {
                        "role": "user",
                        "content": f"RESUME:\n{text}\n\n{EXTRACTION_PROMPT}",
                    }
                ],
            )

            if not response.content:
                raise RuntimeError("Claude returned an empty response.")

            return self._process(response.content[0].text)

        except Exception as exc:
            raise RuntimeError(
                f"Claude resume extraction failed for {source_type}. "
                "Check ANTHROPIC_API_KEY, model name, credits, network, and resume file. "
                f"Details: {exc}"
            ) from exc

    def _clean_model_json(self, raw: str) -> str:
        if raw is None:
            raise RuntimeError("Claude extraction returned None.")

        clean = str(raw).strip()

        if clean.startswith("```"):
            clean = clean.replace("```json", "")
            clean = clean.replace("```", "")
            clean = clean.strip()

        first_brace = clean.find("{")
        last_brace = clean.rfind("}")

        if first_brace != -1 and last_brace != -1 and last_brace > first_brace:
            clean = clean[first_brace:last_brace + 1]

        return clean.strip()

    def _process(self, raw: str) -> Dict[str, Any]:
        """
        Parse Claude JSON response and map it to internal profile schema.

        Invalid JSON raises a clear error showing a small raw preview.
        """
        clean = self._clean_model_json(raw)

        try:
            extracted = json.loads(clean)
        except json.JSONDecodeError as exc:
            preview = clean[:700]
            raise RuntimeError(
                "Claude extraction did not return valid JSON. "
                f"JSON error: {exc}. Raw preview: {preview}"
            ) from exc

        if not isinstance(extracted, dict):
            raise RuntimeError(
                f"Claude extraction must return a JSON object, got {type(extracted).__name__}."
            )

        return self._map(extracted)

    def _safe_list(self, value: Any) -> List[Any]:
        if isinstance(value, list):
            return value
        if value in [None, ""]:
            return []
        return [value]

    def _safe_int(self, value: Any, default: int = 0) -> int:
        try:
            if value in [None, ""]:
                return default
            return int(float(value))
        except Exception:
            return default

    def _safe_float_or_original(self, value: Any) -> Any:
        try:
            if value in [None, ""]:
                return None
            return float(value)
        except Exception:
            return value

    def _normalise_project(self, project: Any) -> Dict[str, Any]:
        if isinstance(project, dict):
            return {
                "title": project.get("title") or project.get("name") or "Untitled Project",
                "description": project.get("description") or "",
                "technologies": self._safe_list(project.get("technologies", [])),
            }

        return {
            "title": str(project),
            "description": "",
            "technologies": [],
        }

    def _map(self, extracted: Dict[str, Any]) -> Dict[str, Any]:
        """Map extraction output to Korgut's internal profile schema."""
        gaps = [str(gap) for gap in self._safe_list(extracted.get("gaps", []))]

        for required_gap in ["budget", "target_disciplines"]:
            if required_gap not in gaps:
                gaps.append(required_gap)

        confidence_notes = extracted.get("confidence_notes")
        notes = "Profile from resume. "

        if confidence_notes:
            notes += str(confidence_notes).strip() + " "

        if gaps:
            notes += f"Aria to collect: {', '.join(gaps)}."

        projects = [
            self._normalise_project(project)
            for project in self._safe_list(extracted.get("projects", []))
        ]

        technical_skills = [
            str(skill).strip()
            for skill in self._safe_list(extracted.get("technical_skills", []))
            if str(skill).strip()
        ]

        disciplines = [
            str(discipline).strip()
            for discipline in self._safe_list(extracted.get("inferred_disciplines", []))
            if str(discipline).strip()
        ]

        profile = {
            "name": extracted.get("name") or "Student",
            "email": extracted.get("email"),
            "institution": extracted.get("undergraduate_institution"),
            "major": extracted.get("undergraduate_major"),
            "program": "MS Computer Science",
            "graduation_year": self._safe_int(extracted.get("graduation_year"), None),
            "gpa": self._safe_float_or_original(extracted.get("gpa")),
            "gpa_scale": extracted.get("gpa_scale") or "4.0",
            "gre_quant": self._safe_int(extracted.get("gre_quant"), None),
            "gre_verbal": self._safe_int(extracted.get("gre_verbal"), None),
            "toefl": self._safe_int(extracted.get("toefl"), None),
            "ielts": self._safe_float_or_original(extracted.get("ielts")),
            "work_months": self._safe_int(extracted.get("work_experience_months"), 0),
            "work_experience_summary": extracted.get("work_experience_summary"),
            "research": extracted.get("research_experience") or "None stated",
            "publications_count": self._safe_int(extracted.get("publications_count"), 0),
            "disciplines": disciplines,
            "skills": technical_skills,
            "technical_skills": technical_skills,
            "projects": projects,
            "budget": None,
            "notes": notes.strip(),
            "source": "resume",
            "verified": False,
            "gaps": gaps,
            "parser_status": "success",
            "parser_engine": "claude",
        }

        return profile

    def print_summary(self, profile: Dict[str, Any]) -> None:
        console.print("\n[bold green]Resume parsed.[/bold green]")

        for label, key in [
            ("Name", "name"),
            ("Institution", "institution"),
            ("Major", "major"),
            ("Program", "program"),
            ("GPA", "gpa"),
            ("Work exp", "work_months"),
            ("Skills", "skills"),
        ]:
            value = profile.get(key)

            if isinstance(value, list):
                value = ", ".join(str(item) for item in value[:5])

            console.print(f"  {label}: {value}")

        gaps = profile.get("gaps", [])

        if gaps:
            console.print(f'  [yellow]Gaps: {", ".join(str(gap) for gap in gaps)}[/yellow]')

        console.print()
